#!/usr/bin/env/ python3
# -*- coding:utf-8 -*-

import os
import re
import time
import codecs
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt, Mm, Cm, RGBColor


def recordClean(file_name, beginTime, endTime, password):
    """
    一、数据清洗
    1）将指定时间之外的消息去除
    2）将表情、图片去除
    二、生成word文档
    三、改变word文档的样式
    """
    t = int(time.time())    # 获取当前时间
    time_struct = time.localtime(t)
    this_time = time.strftime('%Y-%m-%d', time_struct)    # 格式化日期
    print(this_time)

    # 匹配时间
    any_time = re.compile(r'20(\d+?)-(\d+?)-(\d+?)\s+\d*')
    begin_time = re.compile(this_time + r'\s+' + beginTime)
    end_time = re.compile(this_time + r'\s+' + endTime)

    # 匹配图片、表情
    img_pat = re.compile(r'\[图片\]')
    face_pat = re.compile(r'\[表情\]')

    # 打开当前路径指定的文件
    with open(file_name, 'r', encoding='utf-8') as f:
        newline = f.readlines()

    '''
    数据第一次清洗：循环从文件中删除，直到匹配到指定的时间
    '''
    for n in range(len(newline)):
        try:
            if  not re.match(begin_time, newline[0]):
                newline.pop(0)
            else:
                print('匹配到了，这一行是' + newline[0])
                break
        except IndexError as ret:
            pass

    '''
    数据第二次清洗：循环新的newline，将每一行添加到临时content中，直到匹配到结束时间
    '''
    content_tmp = list()    # 临时列表
    for line in newline:
        if not re.match(end_time, line):
            content_tmp.append(line)
        else:
            print('匹配到了结束的' + line)
            break

    with open('content_tmp.txt', 'w', encoding='utf-8') as f1:
        for i in content_tmp:
            f1.write(i)

    '''
    数据第三次清洗：
    1. 时间行
    2. 文字中的表情、图片都不要了
    '''
    content = list()  # 正式内容列表
    with open('content_tmp.txt', 'r', encoding='utf-8') as f2:
        content_line = f2.readlines()

    print(content_line)
    # 由于文字中的图片都已经删掉了，在整理的时候容易不知道在哪儿加图片
    # 文字间的图片可以替换掉
    # 单独一行的图片换成一句话“此处有图片”，样式设置红色
    for line in content_line:
        line.strip()
        if re.findall(r'\n', line):              # 去掉换行,把换行都变成''
            line = re.sub(r'\n', '', line)
        if line == '':                           # 连带着文件中原本的空行，和上一步换行替换的''，都去掉
            continue
        if re.match(r'\[图片\]', line):
            line = re.sub(r'\[图片\]', '此处有图片', line)
        if re.findall(img_pat, line):
            line = re.sub(img_pat, '', line)     # 去掉图片
        if re.findall(face_pat, line):
            line = re.sub(face_pat, '', line)    # 去掉表情
        if re.match(any_time, line):             # 去掉时间
            continue
        if re.match(r'@[^\s]', line):
            line = 'QA: ' + line
        content.append(line)

    # 把临时文件删除
    os.remove('content_tmp.txt')

    # 存入word
    file_word = docx.Document()
    for line in content:
        file_word.add_paragraph(line)
    file_word.save('%s.docx' % file_name[:-4])

    # 设置word样式
    wordFormat(file_name, password)


def wordFormat(file_name, password):
    """
    功能：修改word格式
    :file_name：文件名
    :return:
    """

    # 打开word
    f_word = open('%s.docx'% file_name[:-4], 'rb')
    docu = docx.Document(f_word)
    f_word.close()

    # 字体样式
    font_name_1 = u'微软雅黑'

    # 段落样式
    styles = docu.styles

    '''
    标题的样式
    '''
    title_style = styles.add_style('myTitle', WD_STYLE_TYPE.PARAGRAPH)  # 第一个段样式：myTitle
    title_style.base_style = styles['Normal']  # 继承Normal样式
    title_style.font.name = font_name_1  # 字体
    title_style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    title_style.font.size = Pt(18)  # 字号
    title_style.font.italic = False  # 斜体
    title_style.font.bold = True  # 粗体
    title_style.paragraph_format.line_spacing = Pt(21)  # 行距
    title_style.paragraph_format.space_before = Pt(3)  # 段前间距
    title_style.paragraph_format.space_after = Pt(21)  # 段后间距
    title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    '''
    爬楼密码样式
    '''
    pwd_style = styles.add_style('myPassword', WD_STYLE_TYPE.PARAGRAPH)  # 第二个段样式：myPassword
    pwd_style.base_style = styles['Normal']  # 继承Normal样式
    pwd_style.font.name = font_name_1  # 字体
    pwd_style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    pwd_style.font.size = Pt(14)  # 字号
    pwd_style.font.bold = False  # 粗体
    pwd_style.font.color.rgb = RGBColor(255,0,0)
    pwd_style.paragraph_format.space_after = Pt(21)  # 段后间距
    pwd_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    '''
    正文样式
    '''
    text_style = styles.add_style('myText', WD_STYLE_TYPE.PARAGRAPH) # 正文样式
    text_style.base_style = styles['Normal']  # 继承Normal样式
    text_style.font.name = font_name_1  # 字体
    text_style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    text_style.font.size = Pt(14)  # 字号
    text_style.font.italic = False  # 斜体
    text_style.paragraph_format.line_spacing = 1.5  # 1.5倍行间距
    text_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    text_style.paragraph_format.first_line_indent = Inches(0.5)
    '''
    此处有图片样式
    '''
    img_style = styles.add_style('myImg', WD_STYLE_TYPE.PARAGRAPH)  # 图片样式
    img_style.base_style = styles['Normal']  # 继承Normal样式
    img_style.font.name = font_name_1  # 字体
    img_style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    img_style.font.size = Pt(14)  # 字号
    img_style.font.italic = False  # 斜体
    img_style.font.color.rgb = RGBColor(255,0,0) # 颜色
    img_style.paragraph_format.line_spacing = 1.5  # 1.5倍行间距
    img_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    img_style.paragraph_format.first_line_indent = Inches(0.5)


    # 提取所有段落
    paragraphs = docu.paragraphs

    for i in range(len(paragraphs)):  # 对于每一段
        if i == 0: # 首行上面增加标题和爬楼密码
            prior_paragraph = paragraphs[i].insert_paragraph_before(file_name[:-4])
            prior_paragraph.style = title_style
            pwd_para = paragraphs[i].insert_paragraph_before('爬楼密码：%s' % password)
            pwd_para.style = pwd_style
        if re.match(r'此处有图片', paragraphs[i].text): # 如果有图片，就标红色
            paragraphs[i].style = img_style
        elif re.match(r'QA:\s', paragraphs[i].text): # 如果是@其他人的，当做答疑，标为红色
            paragraphs[i].style = img_style
        else:
            paragraphs[i].style = text_style  # 当前段格式

    docu.save('%s1.docx'% file_name[:-4])


def main():

    # 获取聊天记录文件名、爬楼密码、开始时间、结束时间
    file_name = input('你的聊天记录文件叫什么呐:')
    password = input('爬楼密码：')
    beginTime = input('几点开始的分享呢？格式例子：20：00')
    endTime = input('几点结束的分享呢？格式例子：20:30')

    # file_name = '长投学堂77期42班小白营.txt'
    # password = '12345'
    # beginTime = '20:02'
    # endTime = '20:50'
    recordClean(file_name, beginTime, endTime, password)


if __name__ == '__main__':
    main()